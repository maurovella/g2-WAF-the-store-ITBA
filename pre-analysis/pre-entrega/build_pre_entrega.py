from __future__ import annotations

from html import escape
from pathlib import Path
import subprocess
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "pre-analysis" / "pre-entrega"
OUT_DIR.mkdir(parents=True, exist_ok=True)

TITLE = "Pre-entrega TPE Tema 9 · Protección de Servicios con WAF para The Store"
SUBTITLE = "Tema 9 · Redes de Información · ITBA · 1C 2026"
GROUP = "Grupo 2"
MEMBERS = [
    "Mauro Vella",
    "Enrique Castillo - 68321",
    "Federico Inti Garcia Lauberer - 61374",
]

PROBLEMATICAS = [
    "The Store corre sobre Kind/Kubernetes y expone solo la UI vía `ingress-nginx`; los demás microservicios permanecen detrás de `ClusterIP` y no deberían ser accesibles desde el exterior.",
    "La auditoría pre-WAF realizada contra `http://localhost` confirmó 10 vectores explotables sin autenticación. Los más críticos son SSRF/path traversal vía `/proxy/**`, IDOR de lectura en `/proxy/carts/{customerId}` y exposición de Spring Actuator.",
    "Hoy no hay rate limiting, scanner detection ni headers HTTP de seguridad. Esto deja abierta la puerta a enumeración, scraping, abuso automatizado, fuga de información operativa y reconocimiento previo a ataques más complejos.",
    "El problema es transversal a una arquitectura multi-lenguaje. Corregir individualmente en Java, Go y Node.js requiere cambios distribuidos; un WAF permite aplicar una política unificada en el punto de entrada.",
]

HALLAZGOS = [
    ("H1", "SSRF + path traversal vía `/proxy/**`", "Crítica", "A03 / A10"),
    ("H2", "IDOR de lectura en `/proxy/carts/{customerId}`", "Crítica", "A01"),
    ("H3", "Spring Actuator expuesto (`info`, `health`, `metrics`, `prometheus`)", "Alta", "A05"),
    ("H4", "Sin rate limit ni scanner detection", "Alta", "A04"),
    ("H5", "Headers de seguridad ausentes", "Media", "A05"),
]

DISENIO = [
    "Se propone desplegar `ModSecurity v3` con `OWASP Core Rule Set (CRS)` dentro del `ingress-nginx-controller` ya presente en el cluster Kind.",
    "La solución operará en dos etapas: `DetectionOnly` para tuning inicial y luego `On` para bloqueo efectivo, minimizando falsos positivos antes de la demo final.",
    "Además del CRS base, se incluirán reglas custom para bloquear `/actuator/*` sensibles, negar traversal sobre `/proxy/**`, limitar abuso automatizado y reforzar headers de seguridad.",
    "La principal ventaja es que no requiere cambios funcionales en los microservicios: el cambio queda acotado al Ingress, su ConfigMap y las reglas del WAF.",
]

COMPONENTES = [
    ("ConfigMap de `ingress-nginx`", "Activar `enable-modsecurity`, CRS y `SecRuleEngine` en modo `DetectionOnly`/`On`."),
    ("OWASP CRS", "Cubrir traversal, scanners y payloads genéricos de capa 7 sin tocar el código de negocio."),
    ("Reglas custom", "Bloquear `/actuator/*`, traversal sobre `/proxy/**` y reforzar guardrails contextuales."),
    ("Anotaciones del Ingress", "Aplicar `limit-rps`, `limit-connections` y `more_set_headers` para hardening perimetral."),
]

SCOPE = [
    ("Bloqueo de Actuator sensible", "Requests a `/actuator/info`, `/actuator/metrics` y `/actuator/prometheus` deberán pasar de `200 OK` a `403 Forbidden`."),
    ("Mitigación de SSRF/path traversal vía proxy", "Payloads como `/proxy/catalog/../../actuator/prometheus` deberán quedar bloqueados por CRS 930xxx y/o regla custom específica."),
    ("Detección de scanners conocidos", "User-Agents de herramientas como Nikto o sqlmap deberán quedar registrados y bloqueados por reglas del CRS."),
    ("Rate limiting en el punto de entrada", "Se limitará el caudal por IP desde el Ingress para mitigar abuso automatizado y DoS de baja complejidad."),
    ("Headers de seguridad HTTP", "El Ingress agregará `HSTS`, `X-Frame-Options`, `X-Content-Type-Options`, `Content-Security-Policy`, `Referrer-Policy` y `Permissions-Policy`."),
]

CASOS_USO = [
    "Antes/después de exfiltración de métricas vía `/actuator/prometheus`.",
    "Antes/después de traversal vía `/proxy/catalog/../../actuator/info`.",
    "Request con User-Agent de Nikto antes/después del WAF.",
    "Burst de requests para mostrar ausencia/presencia de rate limiting.",
    "Comparación de headers de respuesta de la home antes/después.",
]

ARCH_BLOCK = """\
Internet / Host (127.0.0.1:80,443)
        │
        ▼
Kind single-node · Ubuntu 22.04 base · HTTP/1.1
        │
        ▼
Namespace ingress-nginx
  ingress-nginx-controller
  + Nginx
  + libmodsecurity v3
  + OWASP CRS
        │
        ▼
Ingress host=localhost path=/  ─────────────►  ui:80
                                              │
                                              ├──► catalog:80
                                              ├──► carts:80
                                              ├──► checkout:80
                                              └──► orders:80
"""

NETWORK_ROWS = [
    ("Host loopback", "127.0.0.1/32", "Usuario accede a The Store por 80/443"),
    ("Docker / nodo Kind", "172.18.0.0/16", "Red del contenedor del nodo"),
    ("Pods", "10.244.0.0/16", "Comunicación entre pods"),
    ("Services", "10.96.0.0/12", "ClusterIP internos"),
]

ALTERNATIVAS = [
    ("Shadow Daemon", "Se descarta porque exige instrumentación más cercana a la aplicación y no encaja naturalmente con una arquitectura Java/Go/Node detrás de un Ingress compartido."),
    ("OpenWAF", "Fue considerado por figurar en el enunciado, pero se evita por menor madurez de ecosistema y menor predictibilidad para una demo académica."),
    ("Coraza", "Es técnicamente interesante y compatible con CRS, pero no figura en el tema asignado; se deja como evolución futura."),
]

VALIDACION = [
    "Usar como baseline la evidencia pre-WAF ya capturada en `pre-analysis/evidencias/`.",
    "Repetir los PoCs seleccionados con `DetectionOnly` y luego con bloqueo activo.",
    "Guardar snapshot post-WAF y comparar resultados pre/post con un diff simple.",
    "Adjuntar logs de ModSecurity para demostrar que el bloqueo provino del WAF.",
]

METRICAS = [
    ("/actuator/prometheus", "200 OK", "403 Forbidden"),
    ("Traversal vía /proxy/**", "explotable", "bloqueado"),
    ("User-Agent Nikto/sqlmap", "permitido", "detectado/bloqueado"),
    ("Rate limit", "no presente", "activo en el Ingress"),
    ("Headers de seguridad", "0/6", "6/6"),
]

RIESGOS = [
    ("Falsos positivos del CRS", "Arrancar en `DetectionOnly`, revisar logs y excluir reglas puntuales solo si afectan tráfico legítimo."),
    ("Bloqueo insuficiente de vectores propios de la app", "Complementar el CRS con reglas custom centradas en `/actuator/*` y `/proxy/**`."),
    ("Dependencia del Ingress como punto único", "Mantener cambios acotados y reproducibles en manifests para poder iterar rápido."),
]

CRONOGRAMA = [
    ("Semana 1", "Validar pre-entrega y congelar alcance definitivo del POC."),
    ("Semana 2", "Desplegar ModSecurity + CRS sobre `ingress-nginx` en el cluster local."),
    ("Semana 3", "Implementar reglas custom y tunear falsos positivos."),
    ("Semana 4", "Generar baseline post-WAF, diff de resultados y material para la demo final."),
]


def md_code_block(text: str) -> str:
    return "```text\n" + text.strip("\n") + "\n```"


def html_inline(text: str) -> str:
    parts = text.split("`")
    rendered: list[str] = []
    for idx, part in enumerate(parts):
        if idx % 2 == 1:
            rendered.append(f"<code>{escape(part)}</code>")
        else:
            rendered.append(escape(part))
    return "".join(rendered)


def members_markdown_block() -> list[str]:
    lines = ["**Integrantes:**  "]
    lines.extend(f"- {member}" for member in MEMBERS)
    return lines


def members_html_block() -> str:
    return "".join(f"<div>{escape(member)}</div>" for member in MEMBERS)


def members_reportlab_block() -> str:
    return "<br/>".join(escape(member) for member in MEMBERS)


def html_list(items: list[str]) -> str:
    return "<ul>" + "".join(f"<li>{html_inline(item)}</li>" for item in items) + "</ul>"


def html_pairs(items: list[tuple[str, str]]) -> str:
    return "<ul>" + "".join(
        f"<li><strong>{escape(a)}:</strong> {html_inline(b)}</li>" for a, b in items
    ) + "</ul>"


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    p.space_after = Pt(0)


def shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def build_markdown() -> str:
    lines: list[str] = [
        f"# {TITLE}",
        "",
        f"**{SUBTITLE}**  ",
        f"**{GROUP}**  ",
        "",
    ]
    lines.extend(members_markdown_block())
    lines.extend(["", "## 1. Problemática y contexto", ""])
    lines.extend(f"- {item}" for item in PROBLEMATICAS)
    lines.extend(["", "### Hallazgos que motivan la solución", "", "| ID | Hallazgo | Severidad | OWASP Top 10 |", "|---|---|---|---|"])
    for row in HALLAZGOS:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} |")

    lines.extend(["", "## 2. Diseño de la solución", ""])
    lines.extend(f"- {item}" for item in DISENIO)
    lines.extend(["", "### Componentes a implementar", ""])
    for title, body in COMPONENTES:
        lines.append(f"- **{title}:** {body}")

    lines.extend(["", "## 3. Scope del POC y casos de uso", "", "### Alcance comprometido", ""])
    for title, body in SCOPE:
        lines.append(f"- **{title}:** {body}")
    lines.extend(["", "### Casos de uso para la demo", ""])
    lines.extend(f"- {item}" for item in CASOS_USO)

    lines.extend(["", "## 4. Diagrama de arquitectura", "", md_code_block(ARCH_BLOCK), "", "| Bloque | CIDR | Rol |", "|---|---|---|"])
    for name, cidr, role in NETWORK_ROWS:
        lines.append(f"| {name} | `{cidr}` | {role} |")
    lines.extend(
        [
            "",
            "Detalles técnicos relevantes:",
            "",
            "- **SO base del nodo:** `kindest/node` sobre Ubuntu 22.04.",
            "- **Protocolos:** HTTP/1.1 externo e interno; DNS interno de Kubernetes.",
            "- **Punto de inserción del WAF:** `ingress-nginx-controller` en el namespace `ingress-nginx`.",
            "- **Cambios esperados:** ConfigMap, reglas ModSecurity/CRS y anotaciones del Ingress principal.",
            "",
            "## 5. Alternativas consideradas",
            "",
        ]
    )
    for title, body in ALTERNATIVAS:
        lines.append(f"- **{title}:** {body}")

    lines.extend(["", "## 6. Plan breve de validación", ""])
    lines.extend(f"- {item}" for item in VALIDACION)
    lines.extend(["", "### Métricas esperadas del antes/después", "", "| Métrica | Pre-WAF | Post-WAF esperado |", "|---|---|---|"])
    for row in METRICAS:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} |")

    lines.extend(["", "## 7. Riesgos y mitigaciones", ""])
    for title, body in RIESGOS:
        lines.append(f"- **{title}:** {body}")

    lines.extend(["", "## 8. Cronograma resumido", ""])
    for semana, tarea in CRONOGRAMA:
        lines.append(f"- **{semana}:** {tarea}")

    lines.extend(
        [
            "",
            "## 9. Criterio de éxito",
            "",
            "Se considerará cumplido el POC si los vectores hoy explotables en el punto de entrada HTTP quedan detectados o bloqueados por el WAF, sin romper el acceso normal a la home, catálogo, carrito y checkout. La entrega final incluirá demo funcional, código en GitHub y how-to de despliegue.",
            "",
        ]
    )
    return "\n".join(lines)


def build_docx(docx_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(SUBTITLE + "\n").italic = True
    meta.add_run(GROUP + "\n").bold = True
    meta.add_run("Integrantes:\n")
    for member in MEMBERS:
        meta.add_run(member + "\n")
    meta.paragraph_format.space_after = Pt(10)

    def add_heading(text: str, level: int = 1) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(12.5 if level == 1 else 11)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)

    def add_bullets(items: list[str]) -> None:
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(0)
            p.add_run(item)

    def add_pair_bullets(items: list[tuple[str, str]]) -> None:
        for left, right in items:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"{left}: ")
            r.bold = True
            p.add_run(right)

    add_heading("1. Problemática y contexto")
    add_bullets(PROBLEMATICAS)

    add_heading("Hallazgos que motivan la solución", level=2)
    table = doc.add_table(rows=1, cols=4)
    style_table(table)
    for idx, text in enumerate(["ID", "Hallazgo", "Severidad", "OWASP"]):
        set_cell_text(table.rows[0].cells[idx], text, bold=True)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    for row in HALLAZGOS:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)

    add_heading("2. Diseño de la solución")
    add_bullets(DISENIO)
    add_heading("Componentes a implementar", level=2)
    add_pair_bullets(COMPONENTES)

    add_heading("3. Scope del POC y casos de uso")
    add_heading("Alcance comprometido", level=2)
    add_pair_bullets(SCOPE)
    add_heading("Casos de uso para la demo", level=2)
    add_bullets(CASOS_USO)

    add_heading("4. Diagrama de arquitectura")
    p = doc.add_paragraph()
    r = p.add_run(ARCH_BLOCK)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    table = doc.add_table(rows=1, cols=3)
    style_table(table)
    for idx, text in enumerate(["Bloque", "CIDR", "Rol"]):
        set_cell_text(table.rows[0].cells[idx], text, bold=True)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    for row in NETWORK_ROWS:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    add_bullets(
        [
            "SO base del nodo: kindest/node sobre Ubuntu 22.04.",
            "Protocolos: HTTP/1.1 externo e interno; DNS interno de Kubernetes.",
            "Punto de inserción del WAF: ingress-nginx-controller en ingress-nginx.",
            "Cambios esperados: ConfigMap, reglas ModSecurity/CRS y anotaciones del Ingress principal.",
        ]
    )

    add_heading("5. Alternativas consideradas")
    add_pair_bullets(ALTERNATIVAS)

    add_heading("6. Plan breve de validación")
    add_bullets(VALIDACION)
    add_heading("Métricas esperadas del antes/después", level=2)
    table = doc.add_table(rows=1, cols=3)
    style_table(table)
    for idx, text in enumerate(["Métrica", "Pre-WAF", "Post-WAF esperado"]):
        set_cell_text(table.rows[0].cells[idx], text, bold=True)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    for row in METRICAS:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)

    add_heading("7. Riesgos y mitigaciones")
    add_pair_bullets(RIESGOS)

    add_heading("8. Cronograma resumido")
    add_pair_bullets(CRONOGRAMA)

    add_heading("9. Criterio de éxito")
    doc.add_paragraph(
        "Se considerará cumplido el POC si los vectores hoy explotables en el punto de entrada HTTP quedan detectados o bloqueados por el WAF, sin romper el acceso normal a la home, catálogo, carrito y checkout. La entrega final incluirá demo funcional, código en GitHub y how-to de despliegue."
    )

    doc.save(docx_path)


def build_html() -> str:
    hallazgos_rows = "".join(
        f"<tr><td>{escape(a)}</td><td>{html_inline(b)}</td><td>{escape(c)}</td><td>{escape(d)}</td></tr>"
        for a, b, c, d in HALLAZGOS
    )
    network_rows = "".join(
        f"<tr><td>{escape(a)}</td><td><code>{escape(b)}</code></td><td>{escape(c)}</td></tr>"
        for a, b, c in NETWORK_ROWS
    )
    metric_rows = "".join(
        f"<tr><td>{html_inline(a)}</td><td>{html_inline(b)}</td><td>{html_inline(c)}</td></tr>"
        for a, b, c in METRICAS
    )
    return f"""<!DOCTYPE html>
<html lang="es">
<head>
  <meta charset="utf-8">
  <title>{escape(TITLE)}</title>
  <style>
    @page {{
      size: A4;
      margin: 15mm 16mm 15mm 16mm;
    }}
    body {{
      font-family: "Segoe UI", Arial, sans-serif;
      font-size: 11.8px;
      line-height: 1.34;
      color: #111827;
      margin: 0;
    }}
    h1 {{
      font-size: 20px;
      text-align: center;
      margin: 0 0 5px 0;
    }}
    .meta {{
      text-align: center;
      margin-bottom: 12px;
      font-size: 11px;
    }}
    h2 {{
      font-size: 13px;
      margin: 10px 0 5px 0;
      border-bottom: 1px solid #cbd5e1;
      padding-bottom: 2px;
    }}
    h3 {{
      font-size: 11.5px;
      margin: 7px 0 3px 0;
    }}
    p {{
      margin: 4px 0;
    }}
    ul {{
      margin: 4px 0 5px 20px;
      padding: 0;
    }}
    li {{
      margin: 0 0 3px 0;
    }}
    code {{
      font-family: Consolas, monospace;
      font-size: 10.5px;
      background: #f3f4f6;
      padding: 0 3px;
      border-radius: 3px;
    }}
    pre {{
      font-family: Consolas, monospace;
      font-size: 10px;
      line-height: 1.25;
      background: #f8fafc;
      border: 1px solid #cbd5e1;
      padding: 8px;
      margin: 5px 0 8px 0;
      white-space: pre;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 6px 0 8px 0;
      font-size: 10.5px;
    }}
    th, td {{
      border: 1px solid #cbd5e1;
      padding: 5px 6px;
      vertical-align: top;
    }}
    th {{
      background: #dbeafe;
      text-align: left;
    }}
    .page {{
      break-after: page;
      page-break-after: always;
    }}
    .page:last-child {{
      break-after: auto;
      page-break-after: auto;
    }}
  </style>
</head>
<body>
  <section class="page">
  <h1>{escape(TITLE)}</h1>
  <div class="meta">
      <div><strong>{escape(SUBTITLE)}</strong></div>
      <div><strong>{escape(GROUP)}</strong></div>
      <div>Integrantes:</div>
      {members_html_block()}
    </div>

    <h2>1. Problemática y contexto</h2>
    {html_list(PROBLEMATICAS)}
    <h3>Hallazgos que motivan la solución</h3>
    <table>
      <thead>
        <tr><th>ID</th><th>Hallazgo</th><th>Severidad</th><th>OWASP Top 10</th></tr>
      </thead>
      <tbody>{hallazgos_rows}</tbody>
    </table>
  </section>

  <section class="page">
    <h2>2. Diseño de la solución</h2>
    {html_list(DISENIO)}
    <h3>Componentes a implementar</h3>
    {html_pairs(COMPONENTES)}

    <h2>3. Scope del POC y casos de uso</h2>
    <h3>Alcance comprometido</h3>
    {html_pairs(SCOPE)}
  </section>

  <section class="page">
    <h2>3. Scope del POC y casos de uso</h2>
    <h3>Casos de uso para la demo</h3>
    {html_list(CASOS_USO)}

    <h2>4. Diagrama de arquitectura</h2>
    <pre>{escape(ARCH_BLOCK)}</pre>
    <table>
      <thead><tr><th>Bloque</th><th>CIDR</th><th>Rol</th></tr></thead>
      <tbody>{network_rows}</tbody>
    </table>
    <ul>
      <li><strong>SO base del nodo:</strong> <code>kindest/node</code> sobre Ubuntu 22.04.</li>
      <li><strong>Protocolos:</strong> HTTP/1.1 externo e interno; DNS interno de Kubernetes.</li>
      <li><strong>Punto de inserción del WAF:</strong> <code>ingress-nginx-controller</code> en el namespace <code>ingress-nginx</code>.</li>
      <li><strong>Cambios esperados:</strong> ConfigMap, reglas ModSecurity/CRS y anotaciones del Ingress principal.</li>
    </ul>
  </section>

  <section class="page">
    <h2>5. Alternativas consideradas</h2>
    {html_pairs(ALTERNATIVAS)}

    <h2>6. Plan breve de validación</h2>
    {html_list(VALIDACION)}
    <h3>Métricas esperadas del antes/después</h3>
    <table>
      <thead><tr><th>Métrica</th><th>Pre-WAF</th><th>Post-WAF esperado</th></tr></thead>
      <tbody>{metric_rows}</tbody>
    </table>

    <h2>7. Riesgos y mitigaciones</h2>
    {html_pairs(RIESGOS)}

    <h2>8. Cronograma resumido</h2>
    {html_pairs(CRONOGRAMA)}

    <h2>9. Criterio de éxito</h2>
    <p>Se considerará cumplido el POC si los vectores hoy explotables en el punto de entrada HTTP quedan detectados o bloqueados por el WAF, sin romper el acceso normal a la home, catálogo, carrito y checkout. La entrega final incluirá demo funcional, código en GitHub y how-to de despliegue.</p>
  </section>
</body>
</html>
"""


def build_pdf_from_edge(html_path: Path, pdf_path: Path) -> str:
    edge = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")
    if not edge.exists():
        return "Microsoft Edge no encontrado; se omitió la exportación a PDF."
    args = [
        str(edge),
        "--headless",
        "--disable-gpu",
        f"--print-to-pdf={pdf_path}",
        html_path.resolve().as_uri(),
    ]
    completed = subprocess.run(args, capture_output=True, text=True)
    if completed.returncode != 0:
        return "Falló la exportación a PDF con Edge."
    return "PDF generado con Edge."


def build_pdf_reportlab(pdf_path: Path) -> tuple[Path, str]:
    regular_font = "Helvetica"
    bold_font = "Helvetica-Bold"
    arial = Path(r"C:\Windows\Fonts\arial.ttf")
    arial_bold = Path(r"C:\Windows\Fonts\arialbd.ttf")
    if arial.exists() and arial_bold.exists():
        pdfmetrics.registerFont(TTFont("ArialCustom", str(arial)))
        pdfmetrics.registerFont(TTFont("ArialCustom-Bold", str(arial_bold)))
        regular_font = "ArialCustom"
        bold_font = "ArialCustom-Bold"

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCustom",
        parent=styles["Title"],
        fontName=bold_font,
        fontSize=19,
        leading=23,
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    meta_style = ParagraphStyle(
        "MetaCustom",
        parent=styles["Normal"],
        fontName=regular_font,
        fontSize=10,
        leading=12,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    cover_meta_style = ParagraphStyle(
        "CoverMetaCustom",
        parent=styles["Normal"],
        fontName=regular_font,
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    cover_label_style = ParagraphStyle(
        "CoverLabelCustom",
        parent=styles["Normal"],
        fontName=bold_font,
        fontSize=11.5,
        leading=15,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=3,
    )
    cover_member_style = ParagraphStyle(
        "CoverMemberCustom",
        parent=styles["Normal"],
        fontName=regular_font,
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=5,
    )
    h_style = ParagraphStyle(
        "HCustom",
        parent=styles["Heading2"],
        fontName=bold_font,
        fontSize=11,
        leading=13,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=4,
        spaceBefore=2,
    )
    subh_style = ParagraphStyle(
        "SubHCustom",
        parent=styles["Heading3"],
        fontName=bold_font,
        fontSize=10.5,
        leading=12.5,
        spaceAfter=3,
        spaceBefore=2,
    )
    body_style = ParagraphStyle(
        "BodyCustom",
        parent=styles["BodyText"],
        fontName=regular_font,
        fontSize=8.9,
        leading=10.7,
        spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "BulletCustom",
        parent=body_style,
        leftIndent=12,
        firstLineIndent=-8,
        bulletIndent=0,
        spaceAfter=2,
    )
    pre_style = ParagraphStyle(
        "PreCustom",
        parent=body_style,
        fontName=regular_font,
        fontSize=8.5,
        leading=10,
        leftIndent=4,
    )

    def bullet(text: str):
        safe = escape(text).replace("\n", "<br/>")
        return Paragraph(f"• {safe}", bullet_style)

    def pair_bullet(left: str, right: str):
        safe_right = escape(right).replace("\n", "<br/>")
        return Paragraph(f"• <b>{escape(left)}:</b> {safe_right}", bullet_style)

    def make_table(data, col_widths):
        table = Table(data, colWidths=col_widths, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
                    ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                    ("FONTNAME", (0, 0), (-1, 0), bold_font),
                    ("FONTNAME", (0, 1), (-1, -1), regular_font),
                    ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                    ("LEADING", (0, 0), (-1, -1), 10),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94a3b8")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        return table

    story = [
        Spacer(1, 42 * mm),
        Paragraph(escape(TITLE), title_style),
        Spacer(1, 4 * mm),
        Spacer(1, 10 * mm),
        Paragraph(escape(GROUP), cover_label_style),
        Spacer(1, 6 * mm),
        Spacer(1, 10 * mm),
        Paragraph("Integrantes", cover_label_style),
    ]
    story.extend(Paragraph(escape(member), cover_member_style) for member in MEMBERS)
    story.extend([Spacer(1, 8 * mm), PageBreak(), Paragraph("1. Problemática y contexto", h_style)])
    story.extend(bullet(item) for item in PROBLEMATICAS)
    story.extend(
        [
            Spacer(1, 3),
            Paragraph("Hallazgos que motivan la solución", subh_style),
            make_table(
                [["ID", "Hallazgo", "Severidad", "OWASP Top 10"], *HALLAZGOS],
                [18 * mm, 88 * mm, 24 * mm, 35 * mm],
            ),
            Paragraph("2. Diseño de la solución", h_style),
        ]
    )
    story.extend(bullet(item) for item in DISENIO)
    story.extend([Spacer(1, 3), Paragraph("Componentes a implementar", subh_style)])
    story.extend(pair_bullet(a, b) for a, b in COMPONENTES)
    story.extend([Spacer(1, 3), Paragraph("3. Scope del POC y casos de uso", h_style), Paragraph("Alcance comprometido", subh_style)])
    story.extend(pair_bullet(a, b) for a, b in SCOPE)
    story.append(PageBreak())
    story.extend(
        [
            Paragraph("3. Scope del POC y casos de uso", h_style),
            Paragraph("Casos de uso para la demo", subh_style),
        ]
    )
    story.extend(bullet(item) for item in CASOS_USO)
    story.extend(
        [
            Spacer(1, 3),
            Paragraph("4. Diagrama de arquitectura", h_style),
            Preformatted(ARCH_BLOCK, pre_style),
            make_table(
                [["Bloque", "CIDR", "Rol"], *NETWORK_ROWS],
                [38 * mm, 34 * mm, 95 * mm],
            ),
            Spacer(1, 3),
        ]
    )
    story.extend(
        bullet(item)
        for item in [
            "SO base del nodo: kindest/node sobre Ubuntu 22.04.",
            "Protocolos: HTTP/1.1 externo e interno; DNS interno de Kubernetes.",
            "Punto de inserción del WAF: ingress-nginx-controller en ingress-nginx.",
            "Cambios esperados: ConfigMap, reglas ModSecurity/CRS y anotaciones del Ingress principal.",
        ]
    )
    story.append(PageBreak())
    story.extend(
        [
            Paragraph("5. Alternativas consideradas", h_style),
        ]
    )
    story.extend(pair_bullet(a, b) for a, b in ALTERNATIVAS)
    story.extend([Spacer(1, 3), Paragraph("6. Plan breve de validación", h_style)])
    story.extend(bullet(item) for item in VALIDACION)
    story.extend(
        [
            Spacer(1, 3),
            Paragraph("Métricas esperadas del antes/después", subh_style),
            make_table(
                [["Métrica", "Pre-WAF", "Post-WAF esperado"], *METRICAS],
                [58 * mm, 36 * mm, 70 * mm],
            ),
            Spacer(1, 3),
            Paragraph("7. Riesgos y mitigaciones", h_style),
        ]
    )
    story.extend(pair_bullet(a, b) for a, b in RIESGOS)
    story.extend([Spacer(1, 3), Paragraph("8. Cronograma resumido", h_style)])
    story.extend(pair_bullet(a, b) for a, b in CRONOGRAMA)
    story.extend(
        [
            Spacer(1, 3),
            Paragraph("9. Criterio de éxito", h_style),
            Paragraph(
                escape(
                    "Se considerará cumplido el POC si los vectores hoy explotables en el punto de entrada HTTP quedan detectados o bloqueados por el WAF, sin romper el acceso normal a la home, catálogo, carrito y checkout. La entrega final incluirá demo funcional, código en GitHub y how-to de despliegue."
                ),
                body_style,
            ),
        ]
    )

    temp_pdf_path = pdf_path.with_name(pdf_path.stem + ".tmp.pdf")

    doc = SimpleDocTemplate(
        str(temp_pdf_path),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    doc.build(story)

    final_pdf_path = pdf_path
    try:
        temp_pdf_path.replace(final_pdf_path)
        return final_pdf_path, "PDF generado con ReportLab."
    except PermissionError:
        fallback_name = (
            f"{pdf_path.stem}-{datetime.now().strftime('%Y%m%d-%H%M%S')}{pdf_path.suffix}"
        )
        final_pdf_path = pdf_path.with_name(fallback_name)
        temp_pdf_path.replace(final_pdf_path)
        return (
            final_pdf_path,
            "El PDF original estaba abierto; se generó una copia con nombre alternativo.",
        )


def main() -> None:
    md_text = build_markdown()
    html_text = build_html()

    md_path = OUT_DIR / "pre-entrega-tpe-waf.md"
    html_path = OUT_DIR / "pre-entrega-tpe-waf.html"
    docx_path = OUT_DIR / "pre-entrega-tpe-waf.docx"
    pdf_path = OUT_DIR / "pre-entrega-tpe-waf-4pag.pdf"

    md_path.write_text(md_text, encoding="utf-8")
    html_path.write_text(html_text, encoding="utf-8")
    build_docx(docx_path)
    generated_pdf_path, pdf_status = build_pdf_reportlab(pdf_path)

    print(f"Markdown: {md_path}")
    print(f"HTML: {html_path}")
    print(f"DOCX: {docx_path}")
    print(f"PDF: {generated_pdf_path if generated_pdf_path.exists() else 'no generado'}")
    print(pdf_status)


if __name__ == "__main__":
    main()
